# -*- coding: utf-8 -*-
"""
Created on Thu Feb 22 18:40:03 2024

@author: Cross Family
"""
# import packages
import time
import pandas as pd
import re
import docx
from docx.shared import RGBColor
import os
from sys import exit
from datetime import datetime
import seaborn as sns
import matplotlib as mpl

#instantiate variables for age calculations
ThisDay = datetime.now().day
ThisMonth = datetime.now().month
ThisYear = datetime.now().year


print("Welcome to the International Personality Item Pool (IPIP)-based 16-factor Personality Inventory.  To begin, please enter your Date of Birth in numerical (YEAR, MM, DD) format")
Year = int(input("Enter YEAR of birth: "))
Month = int(input("Enter Month of birth: "))
Day = int(input ("Enter Day of birth: "))

# calculate current age of examinee
A = int(ThisYear - Year)
B = int(ThisMonth - Month)
C = int(ThisDay - Day)

if C < 0 and B > 0:
    C = int(ThisDay - Day + 30)
    B = int(ThisMonth - Month -1)
elif C < 0 and B < 0:
    A = int(ThisYear - Year - 1)
    B = int(12 + B - 1)
    C = int((ThisDay + 30) - Day)
elif B < 0:
    A = A - 1
    B = abs(B)
elif C < 0 and B == 0:
    C = int((ThisDay + 30) - Day)
    B = int(12 -1)
    A = int(ThisYear - Year - 1)
    

Asub = str(A)
Bsub = str(B)    
Csub = str(C)
print('Your age is ' + Asub + ' years ' + Bsub + ' months ' + Csub + ' days ')
print()

# set assessment items in a dictionary by subscale
assessment = [
    #warmth
    {"question":"I know how to comfort others.", "type":1, "math":"+"},
    {"question":"I enjoy bringing people together.", "type":1, "math":"+"},
    {"question":"I feel others' emotions.", "type":1, "math":"+"},
    {"question":"I don't like to get involved in other people's problems.", "type":1, "math":"-"},
    {"question":"I am not really interested in others.", "type":1, "math":"-"},
    {"question":"I try not to think about the needy.", "type":1, "math":"-"},
    #intellect
    {"question":"I make insightful remarks", "type":2, "math":"+"},
    {"question":"I know the answers to many questions.", "type":2, "math":"+"},
    {"question":"I tend to analyze things.", "type":2, "math":"+"},
    {"question":"I consider myself an average person.", "type":2, "math":"-"},
    {"question":"I get confused easily.", "type":2, "math":"-"},
    {"question":"I know that I am not a special person.", "type":2, "math":"-"},
    #emotional stability
    {"question":"I seldom feel blue.", "type":3, "math":"+"},
    {"question":"I feel comfortable with myself.", "type":3, "math":"+"},
    {"question":"I readily overcome setbacks.", "type":3, "math":"+"},
    {"question":"I have frequent mood swings.", "type":3, "math":"-"},
    {"question":"I often feel blue.", "type":3, "math":"-"},
    {"question":"I dislike myself.", "type":3, "math":"-"},
    #assertiveness
    {"question":"I take charge.", "type":4, "math":"+"},
    {"question":"I want to be in charge.", "type":4, "math":"+"},
    {"question":"I say what I think.", "type":4, "math":"+"},
    {"question":"I wait for others to lead the way.", "type":4, "math":"-"},
    {"question":"I never challenge things.", "type":4, "math":"-"},
    {"question":"I let others make the decisions", "type":4, "math":"-"},
    #gregariousness
    {"question":"I am the life of the party.", "type":5, "math":"+"},
    {"question":"I love large parties.", "type":5, "math":"+"},
    {"question":"I joke around a lot.", "type":5, "math":"+"},
    {"question":"I seldom joke around.", "type":5, "math":"-"},
    {"question":"I don't like crowded events", "type":5, "math":"-"},
    {"question":"I am the last to laugh at a joke.", "type":5, "math":"-"},
   #dutifulness
    {"question":"I believe laws should be strictly enforced.", "type":6, "math":"+"},
    {"question":"I try to follow the rules.", "type":6, "math":"+"},
    {"question":"I believe in one true religion.", "type":6, "math":"+"},
    {"question":"I resist authority.", "type":6, "math":"-"},
    {"question":"I break rules.", "type":6, "math":"-"},
    {"question":"I use swear words.", "type":6, "math":"-"},
    #friendlineess
    {"question":"I feel comfortable around people.", "type":7, "math":"+"},
    {"question":"I talk to a lot of different people at parties.", "type":7, "math":"+"},
    {"question":"I don't mind being the center of attention.", "type":7, "math":"+"},
    {"question":"I find it difficult to approach others.", "type":7, "math":"-"},
    {"question":"I often feel uncomfortable around others.", "type":7, "math":"-"},
    {"question":"I have little to say.", "type":7, "math":"-"},
    #sensitivity
    {"question":"I like to read.", "type":8, "math":"+"},
    {"question":"I enjoy discussing movies and books with others.", "type":8, "math":"+"},
    {"question":"I read a lot.", "type":8, "math":"+"},
    {"question":"I do not enjoy watching dance performances.", "type":8, "math":"-"},
    {"question":"I do not like poetry.", "type":8, "math":"-"},
    {"question":"I dislike works of fiction.", "type":8, "math":"-"},
    #distrust
    {"question":"I find it hard to forgive others.", "type":9, "math":"+"},
    {"question":"I suspect hidden motives in others.", "type":9, "math":"+"},
    {"question":"I am wary of others.", "type":9, "math":"+"},
    {"question":"I trust what other people say.", "type":9, "math":"-"},
    {"question":"I trust others.", "type":9, "math":"-"},
    {"question":"I believe that others have good intentions.", "type":9, "math":"-"},
    #imagination
    {"question":"I do things that others find strange.", "type":10, "math":"+"},
    {"question":"I like to get lost in thought.", "type":10, "math":"+"},
    {"question":"I enjoy wild flights of fantasy.", "type":10, "math":"+"},
    {"question":"I do things by the book.", "type":10, "math":"-"},
    {"question":"I seldom daydream.", "type":10, "math":"-"},
    {"question":"I seldom get lost in thought.", "type":10, "math":"-"},
    #reserve
    {"question":"I reveal little about myself.", "type":11, "math":"+"},
    {"question":"I am hard to get to know.", "type":11, "math":"+"},
    {"question":"I don't talk a lot.", "type":11, "math":"+"},
    {"question":"I am open about myself to others.", "type":11, "math":"-"},
    {"question":"I am open about my feelings.", "type":11, "math":"-"},
    {"question":"I disclose my intimate thoughts.", "type":11, "math":"-"},
    #anxiety
    {"question":"I am afraid that I will do the wrong thing.", "type":12, "math":"+"},
    {"question":"I feel threatened easily.", "type":12, "math":"+"},
    {"question":"I am easily hurt.", "type":12, "math":"+"},
    {"question":"I don't worry about things that have already happened.", "type":12, "math":"-"},
    {"question":"I am not easily bothered by things.", "type":12, "math":"-"},
    {"question":"I don't let others discourage me.", "type":12, "math":"-"},
    #complexity
    {"question":"I believe in the importance of art.", "type": 13, "math":"+"},
    {"question":"I love to think up new ways of doing things.", "type": 13, "math":"+"},
    {"question":"I enjoy hearing new ideas.", "type": 13, "math":"+"},
    {"question":"I avoid philosophical discussions.", "type": 13, "math":"-"},
    {"question":"I rarely look for a deeper meaning in things.", "type": 13, "math":"-"},
    {"question":"I am not interested in theoretical discussions.", "type": 13, "math":"-"},
    #introversion
    {"question":"I want to be left alone.", "type":14, "math":"+"},
    {"question":"I prefer to do things by myself.", "type":14, "math":"+"},
    {"question":"I enjoy spending time by myself.", "type":14, "math":"+"},
    {"question":"I enjoy being part of a group.", "type": 14, "math":"-"},
    {"question":"I enjoy teamwork.", "type": 14, "math":"-"},
    {"question":"I can't do without the company of others.", "type": 14, "math":"-"},
    #orderliness
    {"question":"I want everything to be 'just right'.", "type": 15, "math":"+"},
    {"question":"I get chores done right away.", "type": 15, "math":"+"},
    {"question":"I like order.", "type": 15, "math":"+"},
    {"question":"I am not bothered by messy people.", "type": 15, "math": "-"},
    {"question":"I am not bothered by disorder.", "type": 15, "math": "-"},
    {"question":"I leave a mess in my room.", "type": 15, "math": "-"},
    #emotionality
    {"question":"I get irritated easily.", "type": 16, "math":"+"},
    {"question":"I get angry easily.", "type": 16, "math":"+"},
    {"question":"I am quick to judge others.", "type": 16, "math":"+"},
    {"question":"I am not easily annoyed.", "type": 16, "math": "-"},
    {"question":"I try to forgive and forget.", "type": 16, "math": "-"},
    {"question":"I have a good word for everyone.", "type": 16, "math": "-"},
]
numquestions = len(assessment)

helptext = "Describe yourself as you generally are now, not as you wish to be in the future.\nDescribe yourself as you honestly see yourself, in relation to other people you know of the same sex as you are, and roughly your same age.\nIndicate for each statement which answer best fits as a description of you:\n1. Very Inaccurate\n2. Moderately Inaccurate\n3. Neither Accurate Nor Inaccurate\n4. Moderately Accurate\n5. Very Accurate\n\nOnce you have submitted your answer, you will be asked to type 'y' (or 'Y') to confirm it."

# place answer descriptions in a list
answerdescriptions = [
    "Very Inaccurate",
    "Moderately Inaccurate",
    "Neither Accurate Nor Inaccurate",
    "Moderately Accurate",
    "Very Accurate"
]

# instantiate scores in a list
typeScores = [0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0]
questionnum = 0
print("Welcome to the IPIP 16P Assessment"+str(questionnum)+" Total Questions\n")
print(helptext)
print("\n\nType a number 1-5 to represent your answer. Type 'help' to see this information again during the assessment!")
input("Press ENTER when you are ready to begin...\n")

# Have user enter some data to help save the file uniquely
Name = input("Please enter your initials:  ")

# print question prompts to the examinee, ask for responses, provide the description of the response and get confirmation of answer, then add numerical answer to total for each category 
for questiondata in assessment:
    questionnum = questionnum + 1
    validanswer = False # initial state
    while validanswer == False: # force 1-5 answer, to prevent python error closing script
        print("\nQuestion #" + str(questionnum) + ":")
        answer = input(questiondata['question']+"\n") # ask question
        if answer.isdigit():
            answer = int(answer)
            if answer > 5 or answer < 1:
                print("ERROR: Your answer must be a number 1-5\n")
                validanswer = False
            else:
                print("Your Answer: " + answerdescriptions[answer-1])
                confirm = input("Type Y to confirm your answer, then press ENTER.\n")
                if confirm == "Y" or confirm == "y":
                    if questiondata['math'] == "+":
                        answerMath = answer
                    else:
                        answerMath = 5 - (answer-1)
                    typeScores[int(questiondata['type'])-1] = typeScores[int(questiondata['type'])-1] + answerMath
                    validanswer = True
                else:
                    print("\nPlease answer this question again & confirm it...\nYou can type 'help' for more information!\n")
        elif answer == 'help' or answer == "HELP":
            print("\n\n********************************************\n********************************************\n********************HELP********************\n********************************************\n********************************************")
            print(helptext + "\n\nPlease continue by typing a number 1-5... \n")
        else:
            print("\nYou must answer the question with a number 1-5. Type 'help' for information!")
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
typeinfo = ["Warmth","Intellect","Emotional Stability","Assertiveness", "Gregariousness", "Dutifulness", "Friendliness", "Sensitivity", "Distrust", "Imagination", "Reserve", "Anxiety", "Complexity", "Introversion", "Orderliness", "Emotionality"]
num = 0
x = []
y = []

# print results for each subscale category
for type in typeinfo:
    print(type + ": " + str(typeScores[num]))
    x.append(type)
    y.append(int(typeScores[num]))
    num = num + 1

# create a dataframe to prepare to print to Microsoft Word doc report
df = pd.DataFrame()

df['Personality Factors'] = x
df['Raw Scores'] = y
df['Raw Scores'] = df['Raw Scores'].astype(int)

print(df)
#change strings in y to int data type
y = [int(i) for i in y]
results_dic = dict(zip(x, y))
#find the four top scores in the results dictionary
from heapq import nlargest
N = 4
top_four_factors = nlargest(N, results_dic, key=results_dic.get)
print("Your top four personality factors are: " + str(top_four_factors))

#present results graphically
linegraph = sns.lineplot(df, x="Personality Factors", y="Raw Scores", marker = 'o', markersize = 10, sort=True)
plt.axis('tight')
plt.subplots_adjust(wspace=1,hspace=0.5,left=0.1,top=0.9,right=0.9,bottom=0.2)
plt.grid()

linegraph.set_xticklabels(x, rotation=45,ha="right",rotation_mode='anchor', fontsize = 12)
linegraph.set_title("IPIP-Derived 16DPI Raw Scores")
plt.savefig('linegraph.png')
print("\n\nThank you for taking the IPIP-derived 16 Assessment!  Your printable report is being generated. . .")

doc = docx.Document()

from docx.shared import Pt

from docx.shared import Inches
moment=time.strftime("__%Y-%b-%d",time.localtime())
header = doc.sections[0].header
header.paragraphs[0].text = "IPIP 16-DPI Report: " + Name + moment

h1 = doc.add_heading().add_run("The IPIP-Derived 16-Domain Personality Inventory")
h1.font.name = 'Times New Roman'
h1.font.size = Pt(16)
h1.font.color.rgb = RGBColor(30, 0, 0)
print()
para1 = doc.add_paragraph().add_run("The IPIP-Derived 16-Domain Personality Inventory contains 96 items from the International Personality Item Pool, a public domain repository of items and scales that have been statistically evaluated for reliability. Validity has been established by correlation with other items from instruments with established validity.  The IPIP-derived 16 domain personality inventory items are highly correlated with Cattell's 16 Personality Factor Questionnaire (16PF; Conn & Rieke, 1994). No firm interpretation of scoring is currently available for this assessment. Results of this assessment are designed to provide a relative overview of personality factors; comparison of scores across diverse domains should provide a helpful overview of strong aspects of personality.")
para1.font.name = 'Times New Roman'
para1.font.size = Pt(11)
h2 = doc.add_heading().add_run("IPIP-Derived 16-DPI Raw Scores:", 0)
h2.font.name = 'Times New Roman'
h2.font.size = Pt(14) 
h2.font.color.rgb = RGBColor(30, 0, 0)       
t = doc.add_table(df.shape[0]+1, df.shape[1])

# add the header rows.
for j in range(df.shape[-1]):
    t.cell(0,j).text = df.columns[j]
    
t.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
t.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
t.rows[0].cells[0].paragraphs[0].runs[0].font.underline = True
t.rows[0].cells[1].paragraphs[0].runs[0].font.underline = True

# add the rest of the data frame
for i in range(df.shape[0]):
    for j in range(df.shape[-1]):
        t.cell(i+1,j).text = str(df.values[i,j])
print()
print()
print(linegraph)
doc.add_picture('linegraph.png', width=Inches(6.5), height=Inches(4.5))

#Print out more information about examinee's top four personality factors.
h3 = doc.add_heading().add_run("Your Top Four Personality Factors")
h3.font.name = 'Times New Roman'
h3.font.size = Pt(16)
h3.font.color.rgb = RGBColor(30, 0, 0)

para1 = doc.add_paragraph().add_run("Your top four personality factors are: " + str(top_four_factors[0]) + ", " + str(top_four_factors[1]) +", " +str(top_four_factors[2]) + ", " + "and " + str(top_four_factors[3]) + ".  These factors can be conceptualized on a continuum of their opposites; for instance, those scoring high on the 'distrust' factor would likely score low on a 'trust for others' factor, if such a factor were measured; likewise, those who score high on the 'intellect' factor would likely score lower on a theoretical 'concrete thinking' factor.")
para1.font.name = 'Times New Roman'
para1.font.size = Pt(11)

para2 = doc.add_paragraph().add_run("This assessment and its results are not intended to diagnose or treat any mental health illness or condition. Results should not be interpreted as diagnostic in nature.  Rather, these results are informative, and are intended to increase personal insight of individual characteristics.")
para2.font.name = 'Times New Roman'
para2.font.size = Pt(11)

h4 = doc.add_heading().add_run("References")
h4.font.name = 'Times New Roman'
h4.font.size = Pt(16)
h4.font.color.rgb = RGBColor(30, 0, 0)

para3 = doc.add_paragraph().add_run("Conn, S. R., & Rieke, M. L. (1994). The 16PF fifth edition technical manual. Champaign, IL: Institute for Personality and Ability Testing. \n\nInternational Personality Item Pool: A Scientific Collaboratory for the Development of Advanced Measures of Personality Traits and Other Individual Differences (http://ipip.ori.org/). Internet Web Site.")
para3.font.name = 'Times New Roman'
para3.font.size = Pt(11)

moment=time.strftime("__%Y-%b-%d",time.localtime())
header_name = (Name + moment + '.docx')
doc.save(header_name) 
os.startfile(header_name)           
k=input("press enter to exit")
